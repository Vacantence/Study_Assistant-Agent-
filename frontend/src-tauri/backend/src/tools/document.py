import os
from datetime import datetime

from langchain.tools import tool
from docx import Document

from src.config import Config


def _sanitize_filename(name: str) -> str:
    return "".join(c if c.isalnum() or c in " _-" else "_" for c in name)[:50]


def _generate_markdown(content: str, topic: str) -> str:
    date_str = datetime.now().strftime("%Y-%m-%d")
    filename = f"{date_str}_{_sanitize_filename(topic)}.md"
    filepath = os.path.join(Config.OUTPUT_DIR, filename)
    os.makedirs(Config.OUTPUT_DIR, exist_ok=True)

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"# {topic}\n\n")
        f.write(f"> 生成日期: {date_str}\n\n")
        f.write(content)

    return filepath


def _generate_docx(content: str, topic: str) -> str:
    date_str = datetime.now().strftime("%Y-%m-%d")
    filename = f"{date_str}_{_sanitize_filename(topic)}.docx"
    filepath = os.path.join(Config.OUTPUT_DIR, filename)
    os.makedirs(Config.OUTPUT_DIR, exist_ok=True)

    doc = Document()
    doc.add_heading(topic, 0)
    doc.add_paragraph(f"生成日期: {date_str}")

    for line in content.split("\n"):
        if line.startswith("## "):
            doc.add_heading(line.strip("# "), 2)
        elif line.startswith("### "):
            doc.add_heading(line.strip("# "), 3)
        elif line.startswith("- "):
            doc.add_paragraph(line, style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)

    doc.save(filepath)
    return filepath


@tool
def generate_doc(content: str, topic: str = "学习笔记", fmt: str = "md") -> str:
    """将整理好的学习内容生成文档。fmt 支持 'md'(Markdown) 或 'docx'(Word)。"""
    if fmt == "docx":
        return _generate_docx(content, topic)
    return _generate_markdown(content, topic)
